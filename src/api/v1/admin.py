"""管理后台 API：故事导出、片段删改、用户与 Bot 管理"""
from io import BytesIO
from flask import Blueprint, request, jsonify, current_app, Response
from flask_jwt_extended import jwt_required, get_jwt
from sqlalchemy.orm import Session
import uuid

from src.database import get_db
from src.services.story_service import get_story_by_id
from src.services.branch_service import get_branch_by_id, get_branches_by_story
from src.services.segment_service import get_segments_by_branch, get_segment_by_id


def get_db_session():
    if current_app.config.get('TESTING') and 'TEST_DB' in current_app.config:
        return current_app.config['TEST_DB']
    return next(get_db())


def admin_required(fn):
    """要求 JWT 且 user_type 为 admin"""
    from flask_jwt_extended import verify_jwt_in_request
    def wrapper(*args, **kwargs):
        verify_jwt_in_request()
        if get_jwt().get('user_type') != 'admin':
            return jsonify({'status': 'error', 'error': {'code': 'FORBIDDEN', 'message': '需要管理员权限'}}), 403
        return fn(*args, **kwargs)
    wrapper.__name__ = fn.__name__
    return wrapper


admin_bp = Blueprint('admin', __name__, url_prefix='/admin')


# ---------- 故事导出 ----------
@admin_bp.route('/stories/<story_id>/export', methods=['GET'])
@jwt_required()
@admin_required
def export_story(story_id):
    """
    导出故事为 MD / Word / PDF。
    format=md 返回纯文本 Markdown；word/pdf 后续可接文档生成。
    """
    try:
        story_uuid = uuid.UUID(story_id)
    except ValueError:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '无效的故事ID'}}), 400

    db: Session = get_db_session()
    story = get_story_by_id(db, story_uuid)
    if not story:
        return jsonify({'status': 'error', 'error': {'code': 'NOT_FOUND', 'message': '故事不存在'}}), 404

    fmt = (request.args.get('format') or 'md').lower()
    if fmt not in ('md', 'markdown', 'word', 'docx', 'pdf'):
        fmt = 'md'

    # 主分支：parent_branch 为 None 的第一个分支
    branches, _ = get_branches_by_story(db, story_uuid, limit=500, offset=0)
    main_branch = next((b for b in branches if b.parent_branch is None), branches[0] if branches else None)
    if not main_branch:
        return jsonify({'status': 'error', 'error': {'code': 'NOT_FOUND', 'message': '故事暂无分支'}}), 404

    segments, _ = get_segments_by_branch(db, main_branch.id, limit=5000, offset=0)
    lines = []
    lines.append(f"# {story.title}\n")
    if story.background:
        lines.append(f"{story.background.strip()}\n")
    if story.starter:
        lines.append(f"\n{story.starter.strip()}\n")
    for seg in segments:
        lines.append(seg.content.strip())
        lines.append("")
    body_text = "\n".join(lines).strip()
    safe_title = (story.title or story_id)[:50].replace('/', '_').replace('\\', '_')

    if fmt in ('md', 'markdown'):
        return Response(
            body_text,
            status=200,
            mimetype='text/markdown; charset=utf-8',
            headers={'Content-Disposition': f'attachment; filename="{safe_title}.md"'},
        )

    if fmt in ('word', 'docx'):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            doc = Document()
            doc.add_heading(story.title or 'Untitled', 0)
            if story.background:
                doc.add_paragraph(story.background.strip())
            if story.starter:
                doc.add_paragraph(story.starter.strip())
            for seg in segments:
                doc.add_paragraph(seg.content.strip())
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return Response(
                buf.getvalue(),
                status=200,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                headers={'Content-Disposition': f'attachment; filename="{safe_title}.docx"'},
            )
        except Exception as e:
            return jsonify({'status': 'error', 'error': {'code': 'EXPORT_ERROR', 'message': f'Word 导出失败: {str(e)}'}}), 500

    if fmt == 'pdf':
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            import os
            buf = BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72)
            styles = getSampleStyleSheet()
            font_name = 'Helvetica'
            pdf_font_path = os.environ.get('PDF_FONT_PATH') or current_app.config.get('PDF_FONT_PATH')
            if pdf_font_path and os.path.isfile(pdf_font_path):
                try:
                    pdfmetrics.registerFont(TTFont('CustomCJK', pdf_font_path))
                    font_name = 'CustomCJK'
                except Exception:
                    pass
            title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], fontName=font_name)
            body_style = ParagraphStyle('CustomBody', parent=styles['Normal'], fontName=font_name)
            story_content = []
            story_content.append(Paragraph((story.title or 'Untitled').replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), title_style))
            story_content.append(Spacer(1, 12))
            if story.background:
                for line in (story.background or '').strip().split('\n'):
                    story_content.append(Paragraph(line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style))
            if story.starter:
                for line in (story.starter or '').strip().split('\n'):
                    story_content.append(Paragraph(line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style))
            for seg in segments:
                for line in (seg.content or '').strip().split('\n'):
                    story_content.append(Paragraph(line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style))
            doc.build(story_content)
            buf.seek(0)
            return Response(
                buf.getvalue(),
                status=200,
                mimetype='application/pdf',
                headers={'Content-Disposition': f'attachment; filename="{safe_title}.pdf"'},
            )
        except Exception as e:
            return jsonify({'status': 'error', 'error': {'code': 'EXPORT_ERROR', 'message': f'PDF 导出失败: {str(e)}'}}), 500

    return jsonify({'status': 'error', 'error': {'code': 'UNSUPPORTED', 'message': '仅支持 format=md, word, pdf'}}), 400


# ---------- 片段 PATCH / DELETE ----------
@admin_bp.route('/segments/<segment_id>', methods=['PATCH'])
@jwt_required()
@admin_required
def update_segment(segment_id):
    """修改续写段内容。body: { "content": "新内容" }"""
    try:
        segment_uuid = uuid.UUID(segment_id)
    except ValueError:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '无效的片段ID'}}), 400

    data = request.get_json()
    if not data or 'content' not in data:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '缺少 content'}}), 400

    db: Session = get_db_session()
    segment = get_segment_by_id(db, segment_uuid)
    if not segment:
        return jsonify({'status': 'error', 'error': {'code': 'NOT_FOUND', 'message': '片段不存在'}}), 404

    segment.content = data['content']
    db.commit()
    db.refresh(segment)
    return jsonify({
        'status': 'success',
        'data': {
            'id': str(segment.id),
            'content': segment.content,
            'sequence_order': segment.sequence_order,
        },
    }), 200


@admin_bp.route('/segments/<segment_id>', methods=['DELETE'])
@jwt_required()
@admin_required
def delete_segment(segment_id):
    """删除续写段（物理删除，可能产生 sequence_order 空洞）"""
    try:
        segment_uuid = uuid.UUID(segment_id)
    except ValueError:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '无效的片段ID'}}), 400

    db: Session = get_db_session()
    segment = get_segment_by_id(db, segment_uuid)
    if not segment:
        return jsonify({'status': 'error', 'error': {'code': 'NOT_FOUND', 'message': '片段不存在'}}), 404

    db.delete(segment)
    db.commit()
    return jsonify({'status': 'success', 'data': {'id': segment_id}}), 200


# ---------- 用户列表（内存 users_db，与 auth 一致）----------
@admin_bp.route('/users', methods=['GET'])
@jwt_required()
@admin_required
def list_users():
    """列出所有用户（来自 auth 模块内存存储）"""
    from src.api.v1.auth import users_db
    users = [
        {
            'id': u['id'],
            'username': u.get('username'),
            'email': u.get('email'),
            'user_type': u.get('user_type'),
            'created_at': u.get('created_at'),
        }
        for u in users_db.values()
    ]
    return jsonify({'status': 'success', 'data': {'users': users}}), 200


# ---------- Bot 列表与状态更新 ----------
@admin_bp.route('/bots', methods=['GET'])
@jwt_required()
@admin_required
def list_bots():
    """列出所有 Bot"""
    db: Session = get_db_session()
    from src.models.bot import Bot
    bots = db.query(Bot).order_by(Bot.created_at.desc()).all()
    return jsonify({
        'status': 'success',
        'data': {
            'bots': [
                {
                    'id': str(b.id),
                    'name': b.name,
                    'model': b.model,
                    'status': b.status,
                    'reputation': b.reputation,
                    'created_at': b.created_at.isoformat() if b.created_at else None,
                }
                for b in bots
            ],
        },
    }), 200


@admin_bp.route('/bots/<bot_id>', methods=['PATCH'])
@jwt_required()
@admin_required
def update_bot(bot_id):
    """更新 Bot 状态等。body: { "status": "active"|"suspended" }"""
    try:
        bot_uuid = uuid.UUID(bot_id)
    except ValueError:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '无效的Bot ID'}}), 400

    data = request.get_json() or {}
    db: Session = get_db_session()
    from src.models.bot import Bot
    bot = db.query(Bot).filter(Bot.id == bot_uuid).first()
    if not bot:
        return jsonify({'status': 'error', 'error': {'code': 'NOT_FOUND', 'message': 'Bot不存在'}}), 404

    if 'status' in data and data['status'] in ('active', 'suspended'):
        bot.status = data['status']
    db.commit()
    db.refresh(bot)
    return jsonify({
        'status': 'success',
        'data': {
            'id': str(bot.id),
            'name': bot.name,
            'status': bot.status,
        },
    }), 200


# ========== Bot 分配故事 ==========
@admin_bp.route('/bots/<bot_id>/assign', methods=['POST'])
@jwt_required()
@admin_required
def assign_story_to_bot(bot_id):
    """将故事分配给 Bot
    body: { "story_id": "uuid", "auto_continue": true/false }
    """
    try:
        bot_uuid = uuid.UUID(bot_id)
    except ValueError:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '无效的Bot ID'}}), 400

    data = request.get_json() or {}
    story_id = data.get('story_id')
    if not story_id:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '缺少 story_id'}}), 400
    
    try:
        story_uuid = uuid.UUID(story_id)
    except ValueError:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '无效的 story_id'}}), 400

    db: Session = get_db_session()
    
    from src.models.bot import Bot
    from src.models.agent import AgentStory
    
    # 验证 Bot 存在
    bot = db.query(Bot).filter(Bot.id == bot_uuid).first()
    if not bot:
        return jsonify({'status': 'error', 'error': {'code': 'NOT_FOUND', 'message': 'Bot不存在'}}), 404
    
    # 验证故事存在
    from src.models.story import Story
    story = db.query(Story).filter(Story.id == story_uuid).first()
    if not story:
        return jsonify({'status': 'error', 'error': {'code': 'NOT_FOUND', 'message': '故事不存在'}}), 404
    
    # 检查是否已分配
    existing = db.query(AgentStory).filter(
        AgentStory.agent_id == bot_uuid,
        AgentStory.story_id == story_uuid
    ).first()
    
    if existing:
        # 更新设置
        existing.auto_continue = data.get('auto_continue', True)
        db.commit()
        return jsonify({
            'status': 'success',
            'data': {
                'bot_id': str(bot_id),
                'story_id': str(story_id),
                'auto_continue': existing.auto_continue,
                'message': '已更新分配'
            }
        }), 200
    
    # 创建分配记录
    auto_continue = data.get('auto_continue', True)
    assignment = AgentStory(
        agent_id=bot_uuid,
        story_id=story_uuid,
        auto_continue=auto_continue
    )
    db.add(assignment)
    db.commit()
    
    return jsonify({
        'status': 'success',
        'data': {
            'bot_id': str(bot_id),
            'story_id': str(story_id),
            'auto_continue': auto_continue,
            'message': '分配成功'
        }
    }), 201


@admin_bp.route('/bots/<bot_id>/assignments', methods=['GET'])
@jwt_required()
@admin_required
def get_bot_assignments(bot_id):
    """获取 Bot 分配的所有故事"""
    try:
        bot_uuid = uuid.UUID(bot_id)
    except ValueError:
        return jsonify({'status': 'error', 'error': {'code': 'VALIDATION_ERROR', 'message': '无效的Bot ID'}}), 400

    db: Session = get_db_session()
    from src.models.agent import AgentStory
    from src.models.story import Story
    
    assignments = db.query(AgentStory).filter(AgentStory.agent_id == bot_uuid).all()
    
    from src.models.branch import Branch
    from sqlalchemy import func
    
    result = []
    for a in assignments:
        story = db.query(Story).filter(Story.id == a.story_id).first()
        branches_count = db.query(func.count(Branch.id)).filter(
            Branch.story_id == a.story_id,
            Branch.status == 'active'
        ).scalar() or 0
        
        result.append({
            'story_id': str(a.story_id),
            'story_title': story.title if story else '未知',
            'auto_continue': a.auto_continue,
            'assigned_at': a.created_at.isoformat() if a.created_at else None
        })
    
    return jsonify({
        'status': 'success',
        'data': {
            'bot_id': str(bot_id),
            'assignments': result
        }
    }), 200
